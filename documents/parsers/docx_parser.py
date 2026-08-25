"""DOCX parser — bounded paragraphs/tables; no macros/objects execution."""

from __future__ import annotations

import io
import uuid

from documents.errors import DOCUMENT_MACROS_NOT_ALLOWED, DOCUMENT_PARSE_FAILED, DOCUMENT_TOO_LARGE, DocumentError
from documents.models import DOC_DOCX, ParsedDocument, TableBlock, TextBlock, content_hash_text
from documents.zip_safety import inspect_zip_safety


class DocxDocumentParser:
    parser_id = "docx_v1"
    version = "1.0.0"
    supported_types = (DOC_DOCX,)

    def parse(self, *, document_id: str, data: bytes, filename: str, limits: dict) -> ParsedDocument:
        max_text = int(limits.get("max_text_bytes", 1_000_000))
        zip_info = inspect_zip_safety(data)
        names_l = " ".join(zip_info["names"]).lower()
        if "vbaProject.bin".lower() in names_l or "macro" in names_l:
            raise DocumentError(DOCUMENT_MACROS_NOT_ALLOWED)
        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentError(DOCUMENT_PARSE_FAILED) from exc
        try:
            doc = Document(io.BytesIO(data))
        except Exception as exc:
            raise DocumentError(DOCUMENT_PARSE_FAILED) from exc

        blocks = []
        total = 0
        for i, para in enumerate(doc.paragraphs):
            text = (para.text or "").strip()
            if not text:
                continue
            total += len(text.encode("utf-8"))
            if total > max_text:
                raise DocumentError(DOCUMENT_TOO_LARGE)
            style = getattr(para.style, "name", "") or ""
            section = text if style.startswith("Heading") else None
            blocks.append(
                TextBlock(
                    block_id=str(uuid.uuid4()),
                    ordinal=len(blocks),
                    text=text,
                    content_hash=content_hash_text(text),
                    source_location=f"docx:para:{i}",
                    section=section,
                    metadata_safe={"style": style[:64]},
                )
            )

        tables = []
        for t_i, table in enumerate(doc.tables[:20]):
            rows = []
            for row in table.rows[:100]:
                cells = [((c.text or "").strip())[:500] for c in row.cells[:30]]
                rows.append(tuple(cells))
            if not rows:
                continue
            headers = rows[0]
            body = tuple(rows[1:]) if len(rows) > 1 else ()
            tables.append(
                TableBlock(
                    table_id=str(uuid.uuid4()),
                    ordinal=t_i,
                    name=f"table_{t_i}",
                    rows=body,
                    columns=headers,
                    source_location=f"docx:table:{t_i}",
                )
            )

        return ParsedDocument(
            document_id=document_id,
            text_blocks=tuple(blocks),
            tables=tuple(tables),
            metadata_safe={"filename": filename, "paragraph_count": len(blocks)},
            parser_id=self.parser_id,
            parser_version=self.version,
            title=filename,
        )
