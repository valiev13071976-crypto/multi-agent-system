"""Document generation — DOCX / PDF / TXT from typed data."""

from __future__ import annotations

import hashlib
import io
import uuid

from documents.errors import GENERATION_FAILED, DocumentError
from documents.intelligence.contracts import GeneratedDocument, utc_now
from documents.intelligence.templates import assert_generation_inputs
from documents.models import content_hash_bytes
from documents.platform_models import DocumentTemplate
from security.tenant import normalize_tenant_id


def validate_template_fields(template, data: dict) -> None:
    """Fail closed when required template fields are missing."""
    from documents.errors import GENERATED_DOCUMENT_INVALID, DocumentError

    missing = [f for f in template.required_fields if not str((data or {}).get(f) or "").strip()]
    if missing:
        raise DocumentError(GENERATED_DOCUMENT_INVALID)


def generate_from_template(
    *,
    template,
    data: dict,
    tenant_id: str,
    paragraphs: list[str] | None = None,
    tables: list | None = None,
    headings: list[str] | None = None,
):
    validate_template_fields(template, data)
    title = str(data.get("title") or data.get("subject") or template.template_id)
    body = paragraphs or [str(data.get(k) or "") for k in template.sections if data.get(k)]
    if not body:
        body = [f"{k}: {v}" for k, v in data.items() if k in set(template.required_fields) | set(template.optional_fields)]
    fmt = (template.output_format or "txt").lower()
    if fmt == "docx":
        return generate_docx(
            tenant_id=tenant_id,
            title=title,
            paragraphs=body,
            tables=tables,
            headings=headings,
            template_id=template.template_id,
            template_version=template.version,
        )
    if fmt == "pdf":
        return generate_pdf(
            tenant_id=tenant_id,
            title=title,
            paragraphs=body,
            template_id=template.template_id,
            template_version=template.version,
        )
    return generate_txt(
        tenant_id=tenant_id,
        title=title,
        paragraphs=body,
        template_id=template.template_id,
        template_version=template.version,
    )


def generate_txt(
    *,
    tenant_id: str,
    title: str,
    paragraphs: list[str],
    template_id: str = "txt_basic",
    template_version: str = "1.0.0",
    template: DocumentTemplate | None = None,
    fields: dict | None = None,
) -> GeneratedDocument:
    assert_generation_inputs(title=title, paragraphs=paragraphs, template=template, fields=fields)
    body = title.strip() + "\n\n" + "\n\n".join(p.strip() for p in paragraphs if p)
    data = body.encode("utf-8")
    return GeneratedDocument(
        document_id=str(uuid.uuid4()),
        tenant_id=normalize_tenant_id(tenant_id),
        media_type="text/plain",
        filename=f"{title[:40] or 'document'}.txt",
        content=data,
        template_id=template_id,
        template_version=template_version,
        checksum=content_hash_bytes(data),
        provenance={"generated_at": utc_now().isoformat(), "format": "txt"},
    )


def generate_docx(
    *,
    tenant_id: str,
    title: str,
    paragraphs: list[str],
    tables: list[list[list[str]]] | None = None,
    headings: list[str] | None = None,
    template_id: str = "docx_basic",
    template_version: str = "1.0.0",
    template: DocumentTemplate | None = None,
    fields: dict | None = None,
) -> GeneratedDocument:
    assert_generation_inputs(title=title, paragraphs=paragraphs, template=template, fields=fields)
    try:
        from docx import Document
        from docx.enum.text import WD_BREAK
    except ImportError as exc:
        raise DocumentError(GENERATION_FAILED) from exc
    try:
        doc = Document()
        doc.add_heading(title or "Document", level=1)
        for h in headings or ():
            doc.add_heading(str(h), level=2)
        for p in paragraphs:
            doc.add_paragraph(str(p))
        for table_data in tables or ():
            if not table_data:
                continue
            rows = len(table_data)
            cols = max(len(r) for r in table_data)
            table = doc.add_table(rows=rows, cols=cols)
            for r_i, row in enumerate(table_data):
                for c_i, cell in enumerate(row):
                    table.rows[r_i].cells[c_i].text = str(cell)
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentError(GENERATION_FAILED) from exc
    return GeneratedDocument(
        document_id=str(uuid.uuid4()),
        tenant_id=normalize_tenant_id(tenant_id),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{title[:40] or 'document'}.docx",
        content=data,
        template_id=template_id,
        template_version=template_version,
        checksum=content_hash_bytes(data),
        provenance={"generated_at": utc_now().isoformat(), "format": "docx"},
    )


def generate_pdf(
    *,
    tenant_id: str,
    title: str,
    paragraphs: list[str],
    template_id: str = "pdf_basic",
    template_version: str = "1.0.0",
    template: DocumentTemplate | None = None,
    fields: dict | None = None,
) -> GeneratedDocument:
    """Minimal PDF generation without remote resources (pure PDF operators)."""
    assert_generation_inputs(title=title, paragraphs=paragraphs, template=template, fields=fields)
    try:
        lines = [title] + [p for p in paragraphs if p]
        # Escape PDF string specials
        def _esc(s: str) -> str:
            return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

        content_lines = ["BT", "/F1 12 Tf", "50 750 Td"]
        first = True
        for line in lines[:80]:
            chunk = _esc(str(line)[:120])
            if first:
                content_lines.append(f"({chunk}) Tj")
                first = False
            else:
                content_lines.append(f"0 -16 Td ({chunk}) Tj")
        content_lines.append("ET")
        stream = "\n".join(content_lines).encode("latin-1", errors="replace")
        objects = []
        objects.append(b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n")
        objects.append(b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n")
        objects.append(
            b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>endobj\n"
        )
        objects.append(
            f"4 0 obj<< /Length {len(stream)} >>stream\n".encode("ascii")
            + stream
            + b"\nendstream\nendobj\n"
        )
        objects.append(b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n")
        out = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for obj in objects:
            offsets.append(len(out))
            out.extend(obj)
        xref_pos = len(out)
        out.extend(f"xref\n0 {len(offsets)}\n".encode("ascii"))
        out.extend(b"0000000000 65535 f \n")
        for off in offsets[1:]:
            out.extend(f"{off:010d} 00000 n \n".encode("ascii"))
        out.extend(
            f"trailer<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode(
                "ascii"
            )
        )
        data = bytes(out)
    except Exception as exc:
        raise DocumentError(GENERATION_FAILED) from exc
    return GeneratedDocument(
        document_id=str(uuid.uuid4()),
        tenant_id=normalize_tenant_id(tenant_id),
        media_type="application/pdf",
        filename=f"{title[:40] or 'document'}.pdf",
        content=data,
        template_id=template_id,
        template_version=template_version,
        checksum=content_hash_bytes(data),
        provenance={"generated_at": utc_now().isoformat(), "format": "pdf"},
    )
