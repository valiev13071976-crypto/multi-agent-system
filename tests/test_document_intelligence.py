"""Files & Document Intelligence — contracts, OCR, structured extract, tools."""

from __future__ import annotations

import asyncio
import base64
import io
import unittest
import uuid

from autonomy.capabilities import CAP_FILESYSTEM_READ, CAP_FILESYSTEM_WRITE, CapabilitySet
from autonomy.models import utc_now
from documents.errors import (
    DOCUMENT_MALFORMED,
    DOCUMENT_REQUIRES_OCR,
    DOCUMENT_TYPE_MISMATCH,
    OCR_UNAVAILABLE,
    DocumentError,
)
from documents.intelligence.classify import classify_document_text
from documents.intelligence.compare import compare_structured
from documents.intelligence.contracts import (
    BIZ_CONTRACT,
    BIZ_INVOICE,
    DocumentContent,
    StructuredDocument,
)
from documents.intelligence.convert import convert_document
from documents.intelligence.extraction import extract_structured
from documents.intelligence.generate import generate_docx, generate_pdf, generate_txt
from documents.intelligence.large import LargeDocumentPolicy, build_large_doc_plan
from documents.intelligence.linking import link_documents
from documents.intelligence.ocr import FakeOCRProvider, NullOCRProvider, build_ocr_provider
from documents.intelligence.service import DocumentIntelligenceService, build_document_intelligence
from documents.intelligence.validation import validate_structured
from documents.models import SOURCE_TEST_FIXTURE, DocumentIngestRequest
from documents.parsers import build_default_registry
from documents.parsers.pdf import PdfDocumentParser
from documents.parsers.xml_parser import XmlDocumentParser
from documents.runtime import build_document_runtime
from documents.service import DocumentService
from documents.store import InMemoryDocumentStore
from documents.type_detect import resolve_document_type
from memory.models import SCOPE_PROJECT, MemoryScope
from security.encryption import SENSITIVITY_INTERNAL
from tools.gateway import ToolGateway
from tools.models import TOOL_STATUS_SUCCEEDED, ToolRequest
from tools.platform.bootstrap import register_platform_tools
from tools.registry import ToolRegistry


def _scope(sid="proj-di", tenant="tenant-a"):
    return MemoryScope(scope_type=SCOPE_PROJECT, scope_id=sid, tenant_ref=tenant)


def _caps(*names):
    return CapabilitySet(subject_id="u1", capabilities=names, issued_at=utc_now())


def _run(coro):
    return asyncio.run(coro)


class TypeDetectionTests(unittest.TestCase):
    def test_pdf_magic_wins_over_wrong_extension(self):
        data = b"%PDF-1.4 fake content"
        dtype, media = resolve_document_type(filename="note.txt", data=data)
        self.assertEqual(dtype, "pdf")
        self.assertEqual(media, "application/pdf")

    def test_json_and_xml(self):
        j, _ = resolve_document_type(filename="a.json", data=b'{"a":1}')
        self.assertEqual(j, "json")
        x, _ = resolve_document_type(filename="a.xml", data=b'<?xml version="1.0"?><root/>')
        self.assertEqual(x, "xml")

    def test_png_magic(self):
        png = b"\x89PNG\r\n\x1a\n" + b"\x00" * 20
        dtype, _ = resolve_document_type(filename="scan.bin", data=png)
        self.assertEqual(dtype, "image")

    def test_text_claiming_pdf_mismatch(self):
        with self.assertRaises(DocumentError) as ctx:
            resolve_document_type(filename="x.pdf", data=b"hello plain text")
        self.assertEqual(ctx.exception.reason, DOCUMENT_TYPE_MISMATCH)


class PdfIntelligenceTests(unittest.TestCase):
    def test_malformed_pdf(self):
        with self.assertRaises(DocumentError) as ctx:
            PdfDocumentParser().parse(
                document_id="d1", data=b"not-a-pdf", filename="bad.pdf", limits={}
            )
        self.assertIn(ctx.exception.reason, {DOCUMENT_MALFORMED, "document_parse_failed"})

    def test_blank_requires_ocr(self):
        from pypdf import PdfWriter

        buf = io.BytesIO()
        w = PdfWriter()
        w.add_blank_page(width=72, height=72)
        w.write(buf)
        with self.assertRaises(DocumentError) as ctx:
            PdfDocumentParser().parse(
                document_id="d1", data=buf.getvalue(), filename="blank.pdf", limits={}
            )
        self.assertEqual(ctx.exception.reason, DOCUMENT_REQUIRES_OCR)

    def test_ocr_fallback_with_fake(self):
        from pypdf import PdfWriter

        buf = io.BytesIO()
        w = PdfWriter()
        w.add_blank_page(width=72, height=72)
        w.write(buf)
        intel = DocumentIntelligenceService(ocr_provider=FakeOCRProvider("Scanned invoice text"))
        content = intel.extract_pdf_with_ocr_fallback(
            document_id="d1", data=buf.getvalue(), filename="scan.pdf", tenant_id="t1"
        )
        self.assertIn("invoice", content.text.lower())
        self.assertEqual(content.extraction_method, "ocr")


class DocxSpreadsheetXmlJsonTests(unittest.TestCase):
    def test_docx_order_headings_tables(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx missing")
        doc = Document()
        doc.add_heading("Title One", level=1)
        doc.add_paragraph("Para A")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "SKU"
        table.rows[0].cells[1].text = "Price"
        table.rows[1].cells[0].text = "A1"
        table.rows[1].cells[1].text = "10"
        buf = io.BytesIO()
        doc.save(buf)
        svc = DocumentService(InMemoryDocumentStore())
        row = svc.ingest(
            DocumentIngestRequest(
                scope=_scope(),
                filename="sample.docx",
                content=buf.getvalue(),
                source_type=SOURCE_TEST_FIXTURE,
                source_id="docx1",
                sensitivity=SENSITIVITY_INTERNAL,
            )
        )
        self.assertEqual(row.document_type, "docx")
        chunks = svc.list_chunks(row.document_id, requesting_scope=_scope())
        joined = " ".join(c.content_safe or "" for c in chunks)
        self.assertIn("Title One", joined)
        self.assertIn("Para A", joined)

    def test_xlsx_sheets(self):
        try:
            from openpyxl import Workbook
        except ImportError:
            self.skipTest("openpyxl missing")
        wb = Workbook()
        ws = wb.active
        ws.title = "Prices"
        ws["A1"] = "sku"
        ws["B1"] = "price"
        ws["A2"] = "X1"
        ws["B2"] = 12.5
        ws["C2"] = "=B2*2"
        buf = io.BytesIO()
        wb.save(buf)
        svc = DocumentService(InMemoryDocumentStore())
        row = svc.ingest(
            DocumentIngestRequest(
                scope=_scope(),
                filename="book.xlsx",
                content=buf.getvalue(),
                source_type=SOURCE_TEST_FIXTURE,
                source_id="xlsx1",
                sensitivity=SENSITIVITY_INTERNAL,
            )
        )
        self.assertEqual(row.document_type, "xlsx")
        self.assertGreater(row.chunk_count, 0)

    def test_csv(self):
        data = b"sku,price\nA,1\nB,2\n"
        dtype, _ = resolve_document_type(filename="p.csv", data=data)
        self.assertEqual(dtype, "csv")

    def test_json_ok_and_malformed(self):
        reg = build_default_registry(max_file_bytes=1_000_000)
        parsed = reg.get_parser("json").parse(
            document_id="j1", data=b'{"x":1}', filename="a.json", limits={}
        )
        self.assertTrue(parsed.text_blocks)
        with self.assertRaises(DocumentError) as ctx:
            reg.get_parser("json").parse(
                document_id="j2", data=b"{bad", filename="b.json", limits={}
            )
        self.assertEqual(ctx.exception.reason, DOCUMENT_MALFORMED)

    def test_xml_entity_expansion_denied(self):
        malicious = b"""<?xml version="1.0"?>
<!DOCTYPE foo [<!ENTITY xxe SYSTEM "file:///etc/passwd">]>
<root>&xxe;</root>"""
        with self.assertRaises(DocumentError) as ctx:
            XmlDocumentParser().parse(
                document_id="x1", data=malicious, filename="xxe.xml", limits={}
            )
        self.assertEqual(ctx.exception.reason, DOCUMENT_MALFORMED)


class OcrTests(unittest.TestCase):
    def test_null_provider_explicit(self):
        p = NullOCRProvider()
        self.assertFalse(p.available)
        with self.assertRaises(DocumentError) as ctx:
            p.recognize(b"abc")
        self.assertEqual(ctx.exception.reason, OCR_UNAVAILABLE)

    def test_fake_ocr_image_ingest(self):
        ocr = FakeOCRProvider("Invoice No INV-9\nTotal: 100")
        reg = build_default_registry(max_file_bytes=1_000_000, ocr_provider=ocr)
        png = b"\x89PNG\r\n\x1a\n" + b"\x00" * 64
        parsed = reg.get_parser("image").parse(
            document_id="img1", data=png, filename="scan.png", limits={}
        )
        self.assertIn("INV-9", parsed.text_blocks[0].text)

    def test_build_ocr_disabled(self):
        p = build_ocr_provider({"DOCUMENT_OCR_PROVIDER": "null"})
        self.assertFalse(p.available)


class StructuredExtractionTests(unittest.TestCase):
    def test_contract_invoice_act_waybill_price_list(self):
        contract = DocumentContent(
            document_id="c1",
            text="Contract No CTR-1\nDate: 01.01.2024\nINN: 7707083893\nAmount: 1000\nCurrency: RUB\nSubject: Supply",
        )
        sc = extract_structured(contract, document_type="contract")
        self.assertEqual(sc.document_type, BIZ_CONTRACT)
        self.assertEqual(sc.identifiers.get("contract_number"), "CTR-1")

        inv = DocumentContent(
            document_id="i1",
            text=(
                "Invoice No INV-42\nDate: 02.02.2024\nSupplier: Acme\nBuyer: BuyerCo\n"
                "Subtotal: 100\nVAT: 20\nTotal: 120\nCurrency: RUB"
            ),
        )
        si = extract_structured(inv, document_type="invoice")
        self.assertEqual(si.document_type, BIZ_INVOICE)
        self.assertEqual(si.identifiers.get("invoice_number"), "INV-42")
        self.assertTrue(si.validation_ok)

        act = extract_structured(
            DocumentContent(
                document_id="a1",
                text="Act No ACT-1\nDate: 03.03.2024\nContract No CTR-1\nTotal: 120",
            ),
            document_type="act",
        )
        self.assertEqual(act.document_type, "act")

        way = extract_structured(
            DocumentContent(
                document_id="w1",
                text="Waybill No WB-1\nDate: 04.04.2024\nInvoice No INV-42\nTotal: 120",
            ),
            document_type="waybill",
        )
        self.assertEqual(way.document_type, "waybill")

        pl = extract_structured(
            DocumentContent(
                document_id="p1",
                text="Price list\nSupplier: Acme\nsku,name,price\nA1,Widget,10\nB2,Gadget,20\n",
            ),
            document_type="price_list",
        )
        self.assertEqual(pl.document_type, "price_list")
        self.assertGreaterEqual(len(pl.line_items), 1)

    def test_classify_signals(self):
        biz, conf, signals = classify_document_text("This is an invoice with VAT and НДС")
        self.assertEqual(biz, "invoice")
        self.assertTrue(signals)


class ValidationConfidenceTests(unittest.TestCase):
    def test_totals_inconsistent(self):
        doc = StructuredDocument(
            document_id="v1",
            document_type="invoice",
            schema_version="invoice_v1",
            fields={},
            amounts={"total": 100, "subtotal": 50, "vat_amount": 10},
            identifiers={"invoice_number": "1"},
        )
        vr = validate_structured(doc)
        self.assertFalse(vr.ok)
        self.assertIn("totals_inconsistent", vr.errors)

    def test_confidence_levels_normalized(self):
        c = DocumentContent(document_id="x", text="hi", confidence="bogus")
        self.assertEqual(c.confidence, "medium")


class ComparisonLinkingTests(unittest.TestCase):
    def test_compare_changed_and_line_items(self):
        left = StructuredDocument(
            document_id="l",
            document_type="invoice",
            schema_version="invoice_v1",
            fields={},
            identifiers={"invoice_number": "1"},
            amounts={"total": 100},
            line_items=({"sku": "A", "price": 10},),
        )
        right = StructuredDocument(
            document_id="r",
            document_type="invoice",
            schema_version="invoice_v1",
            fields={},
            identifiers={"invoice_number": "1"},
            amounts={"total": 110},
            line_items=({"sku": "A", "price": 10}, {"sku": "B", "price": 20}),
        )
        result = compare_structured(left, right)
        self.assertFalse(result.unchanged)
        self.assertTrue(any(c["field"] == "total" for c in result.changed_fields))
        self.assertTrue(any(d["op"] == "added" for d in result.table_differences))

    def test_unchanged(self):
        doc = StructuredDocument(
            document_id="u",
            document_type="contract",
            schema_version="contract_v1",
            fields={"subject": "x"},
            identifiers={"contract_number": "C1"},
        )
        r = compare_structured(doc, doc)
        self.assertTrue(r.unchanged)

    def test_link_invoice_act(self):
        inv = StructuredDocument(
            document_id="i",
            document_type="invoice",
            schema_version="invoice_v1",
            fields={},
            identifiers={"invoice_number": "INV-1"},
            amounts={"total": 50},
        )
        act = StructuredDocument(
            document_id="a",
            document_type="act",
            schema_version="act_v1",
            fields={},
            identifiers={"act_number": "A1", "related_contract": "C1"},
            amounts={"total": 50},
        )
        link = link_documents(inv, act)
        self.assertIn("same_total", link.evidence)


class GenerationConversionTests(unittest.TestCase):
    def test_generate_txt_pdf_docx(self):
        t = generate_txt(tenant_id="t1", title="Hello", paragraphs=["Body"])
        self.assertEqual(t.tenant_id, "t1")
        self.assertTrue(t.content.startswith(b"Hello"))
        p = generate_pdf(tenant_id="t1", title="PDF", paragraphs=["Line"])
        self.assertTrue(p.content.startswith(b"%PDF"))
        try:
            d = generate_docx(tenant_id="t1", title="Doc", paragraphs=["P"], headings=["H2"])
            self.assertTrue(d.content[:2] == b"PK")
        except DocumentError:
            self.skipTest("python-docx missing")

    def test_convert_builtin_and_unavailable(self):
        out = convert_document(
            tenant_id="t1",
            source_media_type="text/plain",
            target_format="pdf",
            text="Hello convert",
        )
        self.assertTrue(out.content.startswith(b"%PDF"))
        with self.assertRaises(DocumentError) as ctx:
            convert_document(
                tenant_id="t1",
                source_media_type="text/plain",
                target_format="pdf",
                text="x",
                backend="libreoffice",
            )
        self.assertEqual(ctx.exception.reason, "conversion_unavailable")


class LargeDocTests(unittest.TestCase):
    def test_threshold_plans_batches(self):
        policy = LargeDocumentPolicy(max_sync_pages=5)
        self.assertTrue(policy.requires_async(page_count=20))
        plan = build_large_doc_plan(document_id="d", tenant_id="t", page_count=25, batch_size=10)
        self.assertEqual(plan["batch_count"], 3)
        self.assertEqual(plan["workflow_type"], "document.large_extract")
        intel = DocumentIntelligenceService()
        p = intel.plan_large_extraction(
            document_id="d", tenant_id="t", page_count=100, size_bytes=10
        )
        self.assertTrue(p["async"])


class SecurityTenantTests(unittest.TestCase):
    def test_tenant_isolation_on_extract(self):
        store = InMemoryDocumentStore()
        svc = DocumentService(store)
        row = svc.ingest(
            DocumentIngestRequest(
                scope=_scope(tenant="tenant-a"),
                filename="a.txt",
                content=b"secret for A",
                source_type=SOURCE_TEST_FIXTURE,
                source_id="sec1",
                sensitivity=SENSITIVITY_INTERNAL,
            )
        )
        intel = DocumentIntelligenceService(svc, ocr_provider=NullOCRProvider())
        with self.assertRaises(DocumentError) as ctx:
            intel.extract_content(row.document_id, tenant_id="tenant-b")
        self.assertEqual(ctx.exception.reason, "document_access_denied")

    def test_observability_no_raw_content(self):
        events = []

        class Obs:
            def emit(self, event_type, **kwargs):
                events.append((event_type, kwargs))

        svc = DocumentService(InMemoryDocumentStore(), observability=Obs())
        svc.ingest(
            DocumentIngestRequest(
                scope=_scope(),
                filename="log.txt",
                content=b"SHOULD_NOT_APPEAR_IN_LOGS_XYZ",
                source_type=SOURCE_TEST_FIXTURE,
                source_id="log1",
                sensitivity=SENSITIVITY_INTERNAL,
            )
        )
        blob = str(events)
        self.assertNotIn("SHOULD_NOT_APPEAR_IN_LOGS_XYZ", blob)


class ToolAndAcquisitionTests(unittest.TestCase):
    def test_tools_detect_compare_generate(self):
        intel = build_document_intelligence(
            env={"DOCUMENT_OCR_PROVIDER": "fake"},
            ocr_provider=FakeOCRProvider("ocr text"),
        )
        reg = ToolRegistry()
        register_platform_tools(reg, document_intelligence=intel, env={})
        gw = ToolGateway(registry=reg, register_search=False)
        for tool_id in (
            "document.detect",
            "document.structured_extract",
            "document.compare",
            "document.generate",
            "document.convert",
            "document.ocr",
        ):
            self.assertIsNotNone(reg.get(tool_id))

        det = _run(
            gw.invoke(
                ToolRequest(
                    request_id=str(uuid.uuid4()),
                    workflow_id="wf",
                    task_id="t",
                    tool_id="document.detect",
                    operation="detect",
                    tenant_id="t1",
                    arguments={
                        "filename": "x.txt",
                        "content_b64": base64.b64encode(b"%PDF-1.4").decode(),
                    },
                    requested_capabilities=(CAP_FILESYSTEM_READ,),
                ),
                capabilities=_caps(CAP_FILESYSTEM_READ),
            )
        )
        self.assertEqual(det.status, TOOL_STATUS_SUCCEEDED)
        self.assertEqual(det.data["document_type"], "pdf")

        gen = _run(
            gw.invoke(
                ToolRequest(
                    request_id=str(uuid.uuid4()),
                    workflow_id="wf",
                    task_id="t",
                    tool_id="document.generate",
                    operation="generate",
                    tenant_id="tenant-gen",
                    arguments={
                        "format": "pdf",
                        "title": "Gen",
                        "paragraphs": ["Hello"],
                    },
                    requested_capabilities=(CAP_FILESYSTEM_WRITE,),
                ),
                capabilities=_caps(CAP_FILESYSTEM_WRITE),
            )
        )
        self.assertEqual(gen.status, TOOL_STATUS_SUCCEEDED)
        self.assertEqual(gen.data["tenant_id"], "tenant-gen")
        self.assertTrue(base64.b64decode(gen.data["content_b64"]).startswith(b"%PDF"))

        cmp_res = _run(
            gw.invoke(
                ToolRequest(
                    request_id=str(uuid.uuid4()),
                    workflow_id="wf",
                    task_id="t",
                    tool_id="document.compare",
                    operation="compare",
                    tenant_id="t1",
                    arguments={
                        "left_text": "Invoice No 1\nTotal: 10",
                        "right_text": "Invoice No 1\nTotal: 20",
                        "left_document_type": "invoice",
                        "right_document_type": "invoice",
                    },
                    requested_capabilities=(CAP_FILESYSTEM_READ,),
                ),
                capabilities=_caps(CAP_FILESYSTEM_READ),
            )
        )
        self.assertEqual(cmp_res.status, TOOL_STATUS_SUCCEEDED)
        self.assertFalse(cmp_res.data["unchanged"])

    def test_acquisition_price_list_bridge(self):
        from datetime import datetime, timezone

        from acquisition.models import RawArtifact, RECORD_SUPPLIER_ITEM
        from acquisition.parsers.document_bridge import DocumentBridgeParser
        from documents.models import content_hash_text

        intel = DocumentIntelligenceService()
        structured = extract_structured(
            DocumentContent(
                document_id="pl1",
                text="Price list\nSupplier: X\nsku,name,price\nS1,Item,9.5\n",
            ),
            document_type="price_list",
        )
        text, ctype, meta = intel.to_acquisition_artifact_text(structured)
        art = RawArtifact(
            artifact_id=str(uuid.uuid4()),
            source_id="doc",
            tenant_id="t1",
            content_type=ctype,
            fetched_at=datetime.now(timezone.utc),
            content_text=text,
            checksum=content_hash_text(text),
            document_id="pl1",
            metadata=meta,
        )
        records = DocumentBridgeParser().parse(art)
        self.assertTrue(records)
        self.assertEqual(records[0].record_type, RECORD_SUPPLIER_ITEM)


class RuntimeWiringTests(unittest.TestCase):
    def test_runtime_exposes_intelligence(self):
        rt = build_document_runtime(env={"DOCUMENT_OCR_PROVIDER": "null", "DOCUMENTS_ENABLED": "true"})
        self.assertIsNotNone(rt)
        self.assertIsNotNone(rt.intelligence)
        health = rt.health()
        self.assertIn("ocr_available", health)
        self.assertFalse(health["ocr_available"])


if __name__ == "__main__":
    unittest.main()
